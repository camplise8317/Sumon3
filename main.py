import streamlit as st
import pandas as pd
import google.generativeai as genai
import PyPDF2
import docx
import re
import io 
import openai

# --- Configuración de la API de Gemini y OpenAI ---
st.sidebar.header("Configuración de API Keys")
gemini_api_key = st.sidebar.text_input("API Key de Google Gemini", type="password", help="Obtén tu clave en https://aistudio.google.com/app/apikey")
openai_api_key = st.sidebar.text_input("API Key de OpenAI (para modelos GPT)", type="password", help="Obtén tu clave en https://platform.openai.com/account/api-keys")

# Inicialización condicional de Gemini y OpenAI
gemini_config_ok = False
openai_config_ok = False

if gemini_api_key:
    try:
        genai.configure(api_key=gemini_api_key)
        gemini_config_ok = True
        st.sidebar.success("API Key de Gemini configurada.")
    except Exception as e:
        st.sidebar.error(f"Error al configurar la API Key de Gemini: {e}")
else:
    st.sidebar.warning("Por favor, ingresa tu API Key de Gemini para usar modelos Gemini.")

if openai_api_key:
    openai.api_key = openai_api_key
    openai_config_ok = True
    st.sidebar.success("API Key de OpenAI configurada.")
else:
    st.sidebar.warning("Por favor, ingresa tu API Key de OpenAI para usar modelos GPT.")

# --- Funciones de Lectura de Archivos (Adaptadas para Streamlit Uploader) ---
@st.cache_data 
def leer_excel_cargado(uploaded_file):
    """
    Lee un archivo Excel cargado por Streamlit y lo carga en un DataFrame de pandas.
    """
    if uploaded_file is not None:
        try:
            df = pd.read_excel(uploaded_file)
            st.sidebar.success(f"Archivo Excel '{uploaded_file.name}' cargado exitosamente.")
            return df
        except Exception as e:
            st.sidebar.error(f"Ocurrió un error al leer el archivo Excel: {e}")
            return None
    return None

@st.cache_data 
def leer_pdf_cargado(uploaded_file):
    """
    Lee el texto de un archivo PDF cargado por Streamlit.
    """
    if uploaded_file is not None:
        try:
            texto_pdf = ""
            reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            for page_num in range(len(reader.pages)):
                texto_pdf += reader.pages[page_num].extract_text()
            st.sidebar.success(f"Archivo PDF '{uploaded_file.name}' leído exitosamente.")
            return texto_pdf
        except Exception as e:
            st.sidebar.error(f"Ocurrió un error al leer el archivo PDF: {e}")
            return ""
    return ""

# --- Función para obtener la descripción de la taxonomía de Bloom ---
def get_descripcion_bloom(proceso_cognitivo_elegido):
    descripcion_bloom_map = {
        "RECORDAR": "Recuperar información relevante desde la memoria de largo plazo.",
        "COMPRENDER": "Construir significado a partir de información mediante interpretación, resumen, explicación u otras tareas.",
        "APLICAR": "Usar procedimientos en situaciones conocidas o nuevas.",
        "ANALIZAR": "Descomponer información y examinar relaciones entre partes.",
        "EVALUAR": "Emitir juicios basados en criterios para valorar ideas o soluciones.",
        "CREAR": "Generar nuevas ideas, productos o formas de reorganizar información."
    }
    return descripcion_bloom_map.get(str(proceso_cognitivo_elegido).upper(), "Descripción no disponible para este proceso cognitivo.")

# --- Función para generar texto con Gemini o GPT ---
def generar_texto_con_llm(model_type, model_name, prompt):
    if model_type == "Gemini":
        if not gemini_config_ok:
            st.error("API Key de Gemini no configurada. No se puede generar texto con Gemini.")
            return None
        try:
            modelo = genai.GenerativeModel(model_name)
            response = modelo.generate_content(prompt)
            return response.text
        except Exception as e:
            st.error(f"Error al llamar a Gemini: {e}")
            return None
    elif model_type == "GPT":
        if not openai_config_ok:
            st.error("API Key de OpenAI no configurada. No se puede generar texto con GPT.")
            return None
        try:
            client = openai.OpenAI(api_key=openai.api_key)
            response = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000 
            )
            return response.choices[0].message.content
        except Exception as e:
            st.error(f"Error al llamar a OpenAI: {e}")
            return None
    return None

# --- Función para auditar el ítem generado ---
def auditar_item_con_llm(model_type, model_name, item_generado, grado, area, asignatura, estacion, 
                         proceso_cognitivo, nanohabilidad, microhabilidad, 
                         competencia_nanohabilidad, contexto_educativo, manual_reglas_texto="", descripcion_bloom="", grafico_necesario="", descripcion_grafico="", prompt_auditor_adicional=""):
    """
    Audita un ítem generado para verificar su cumplimiento con criterios específicos.
    """
    auditoria_prompt = f"""
    Eres un experto en validación de ítems educativos, especializado en pruebas tipo ICFES y las directrices del equipo IMPROVE.
    Tu tarea es AUDITAR RIGUROSAMENTE el siguiente ítem generado por un modelo de lenguaje.

    Debes verificar que el ítem cumpla con TODOS los siguientes criterios, prestando especial atención a la alineación con los parámetros proporcionados y a las reglas de formato y contenido.

    --- CRITERIOS DE AUDITORÍA ---
    1.  **Formato del Enunciado:** ¿El enunciado está formulado como pregunta clara y directa, sin ambigüedades ni errores?
    2.  **Número de Opciones:** ¿Hay exactamente 4 opciones (A, B, C, D)?
    3.  **Respuesta Correcta Indicada:** ¿La sección 'RESPUESTA CORRECTA:' está claramente indicada y coincide con una de las opciones?
    4.  **Diseño de Justificaciones:**
        * ¿Hay justificaciones bien diferenciadas para CADA opción (A, B, C, D)?
        * ¿La justificación de la opción **correcta** explica el razonamiento, procedimiento o estrategia relevante (NO por descarte)?
        * ¿Las justificaciones de las opciones **incorrectas** están redactadas siguiendo el formato: “El estudiante podría escoger la opción X porque… Sin embargo, esto es incorrecto porque…”?
    5.  **Estilo y Restricciones:** ¿No se usan negaciones mal redactadas, nombres reales, marcas, lugares reales, datos personales o frases vagas como “ninguna de las anteriores” o “todas las anteriores”?
    6.  **Alineación del Contenido:** ¿El ítem (contexto, enunciado, opciones) está alineado EXCLUSIVAMENTE con los siguientes elementos temáticos y cognitivos?
        * Grado: {grado}
        * Área: {area}
        * Asignatura: {asignatura}
        * Estación o unidad temática: {estacion}
        * Proceso Cognitivo (Taxonomía de Bloom): {proceso_cognitivo} (su descripción es "{descripcion_bloom}")
        * Nanohabilidad (foco principal): {nanohabilidad}
        * Microhabilidad (evidencia de aprendizaje): {microhabilidad}
        * Competencia (asociada a Nanohabilidad): {competencia_nanohabilidad}
        * Nivel educativo del estudiante: {contexto_educativo}
    7.  **Gráfico (si aplica):** Si el ítem indica que requiere un gráfico, ¿la descripción del gráfico es clara, detallada y funcional para su futura creación?
        * Gráfico Necesario: {grafico_necesario}
        * Descripción del Gráfico: {descripcion_grafico if grafico_necesario == 'SÍ' else 'N/A'}

    --- MANUAL DE REGLAS ADICIONAL ---
    Las siguientes reglas son de suma importancia para la calidad y pertinencia del ítem. Debes asegurar que el ítem cumple con todas ellas.
    {manual_reglas_texto}
    -----------------------------------

    --- INSTRUCCIONES ADICIONALES PARA LA AUDITORÍA ---
    {prompt_auditor_adicional if prompt_auditor_adicional else "No se proporcionaron instrucciones adicionales para la auditoría."}
    ---------------------------------------------------

    ÍTEM A AUDITAR:
    --------------------
    {item_generado}
    --------------------

    Devuelve tu auditoría con este formato estructurado:

    VALIDACIÓN DE CRITERIOS:
    - Formato del Enunciado: [✅ / ❌] + Comentario (si ❌)
    - Número de Opciones (4): [✅ / ❌]
    - Respuesta Correcta Indicada: [✅ / ❌]
    - Diseño de Justificaciones: [✅ / ⚠️ / ❌] + Observaciones (si ⚠️/❌)
    - Estilo y Restricciones: [✅ / ⚠️ / ❌] + Observaciones (si ⚠️/❌)
    - Alineación del Contenido: [✅ / ❌] + Comentario (si ❌)
    - Gráfico (si aplica): [✅ / ⚠️ / ❌] + Observaciones (si ⚠️/❌)

    DICTAMEN FINAL:
    [✅ CUMPLE TOTALMENTE / ⚠️ CUMPLE PARCIALMENTE / ❌ RECHAZADO]

    OBSERVACIONES FINALES:
    [Explica de forma concisa qué aspectos necesitan mejora, si el dictamen no es ✅. Si es ✅, puedes indicar "El ítem cumple con todos los criterios."]
    """
    return generar_texto_con_llm(model_type, model_name, auditoria_prompt), auditoria_prompt # Retorna también el prompt de auditoría

# --- Función para generar preguntas usando el modelo de generación seleccionado ---
def generar_pregunta_con_seleccion(gen_model_type, gen_model_name, audit_model_type, audit_model_name, 
                                     fila_datos, criterios_generacion, manual_reglas_texto="", 
                                     informacion_adicional_usuario="", 
                                     prompt_bloom_adicional="", prompt_construccion_adicional="", prompt_especifico_adicional="", 
                                     prompt_auditor_adicional="",
                                     contexto_general_estacion="", feedback_usuario=""): # Añade el feedback del usuario
    """
    Genera una pregunta educativa de opción múltiple usando el modelo de generación seleccionado
    y la itera para refinarla si la auditoría lo requiere.
    """
    tipo_pregunta = criterios_generacion.get("tipo_pregunta", "opción múltiple con 4 opciones") 
    dificultad = criterios_generacion.get("dificultad", "media")
    contexto_educativo = criterios_generacion.get("contexto_educativo", "general")
    formato_justificacion = criterios_generacion.get("formato_justificacion", """
        • Justificación correcta: debe explicar el razonamiento o proceso cognitivo (NO por descarte).
        • Justificaciones incorrectas: deben redactarse como: “El estudiante podría escoger la opción X porque… Sin embargo, esto es incorrecto porque…”
    """)
    
    grado_elegido = fila_datos.get('GRADO', 'no especificado')
    area_elegida = fila_datos.get('ÁREA', 'no especificada')
    asignatura_elegida = fila_datos.get('ASIGNATURA', 'no especificada')
    estacion_elegida = fila_datos.get('ESTACIÓN', 'no especificada')
    proceso_cognitivo_elegido = fila_datos.get('PROCESO COGNITIVO', 'no especificado')
    nanohabilidad_elegida = fila_datos.get('NANOHABILIDAD', 'no especificada')
    microhabilidad_elegida = fila_datos.get('MICROHABILIDAD', 'no especificada')
    competencia_nanohabilidad_elegida = fila_datos.get('COMPETENCIA NANOHABILIDAD', 'no especificada')

    dato_para_pregunta_foco = nanohabilidad_elegida
    descripcion_bloom = get_descripcion_bloom(proceso_cognitivo_elegido)

    current_item_text = ""
    auditoria_status = "❌ RECHAZADO" # Estado inicial
    audit_observations = "" # Observaciones para el refinamiento
    max_refinement_attempts = 5 # Número máximo de intentos de refinamiento
    attempt = 0
    grafico_necesario = "NO" # Valor por defecto
    descripcion_grafico = "" # Valor por defecto

    # Almacenar detalles de clasificación para el ítem
    classification_details = {
        "Grado": grado_elegido,
        "Área": area_elegida,
        "Asignatura": asignatura_elegida,
        "Estación": estacion_elegida,
        "Proceso Cognitivo": proceso_cognitivo_elegido,
        "Nanohabilidad": nanohabilidad_elegida,
        "Microhabilidad": microhabilidad_elegida,
        "Competencia Nanohabilidad": competencia_nanohabilidad_elegida
    }

    item_final_data = None # Variable para guardar el ítem final (aprobado o la última versión auditada)
    full_generation_prompt = "" # Variable para almacenar el prompt completo del generador
    full_auditor_prompt = "" # Variable para almacenar el prompt completo del auditor

    # Añade el feedback del usuario al prompt principal del generador
    prompt_con_feedback = ""
    if feedback_usuario:
        prompt_con_feedback = f"--- RETROALIMENTACIÓN DE USUARIO PARA REFINAMIENTO ---\n{feedback_usuario}\n---------------------------------------------------"

    while auditoria_status != "✅ CUMPLE TOTALMENTE" and attempt < max_refinement_attempts:
        attempt += 1
        # st.info(f"--- Generando/Refinando Ítem (Intento {attempt}/{max_refinement_attempts}) ---") # Comentado para no saturar si son muchos ítems

        # Construcción del prompt para el GENERADOR
        prompt_content_for_llm = f"""
        Eres un diseñador experto en ítems de evaluación educativa, especializado en pruebas tipo ICFES u otras de alta calidad técnica.

        Tu tarea es construir un ítem de {tipo_pregunta} con una única respuesta correcta, cumpliendo rigurosamente las reglas de construcción de ítems y alineado con el marco cognitivo de la Taxonomía de Bloom.

        --- CONTEXTO Y PARÁMETROS DEL ÍTEM ---
        - Grado: {grado_elegido}
        - Área: {area_elegida}
        - Asignatura: {asignatura_elegida}
        - Estación o unidad temática: {estacion_elegida}
        - Proceso cognitivo (Taxonomía de Bloom): {proceso_cognitivo_elegido}
        - Descripción del proceso cognitivo:
          "{descripcion_bloom}"
        
        --- PROMPT ADICIONAL: TAXONOMÍA DE BLOOM / PROCESOS COGNITIVOS ---
        {prompt_bloom_adicional if prompt_bloom_adicional else "No se proporcionaron prompts adicionales específicos para taxonomía de Bloom."}
        ------------------------------------------------------------------

        - Nanohabilidad (foco principal del ítem): {nanohabilidad_elegida}
        - Nivel educativo esperado del estudiante: {contexto_educativo}
        - Nivel de dificultad deseado: {dificultad}

        --- CONTEXTO GENERAL DE LA ESTACIÓN (si aplica) ---
        {f"Considera este contexto general para todos los ítems de esta estación: {contexto_general_estacion}" if contexto_general_estacion else "Este ítem debe generar su propio contexto individual, o no se ha definido un contexto general para la estación."}
        ----------------------------------------------------

        --- INSTRUCCIONES PARA LA CONSTRUCCIÓN DEL ÍTEM ---
        CONTEXTO DEL ÍTEM:
        - Incluye una situación contextualizada, relevante y plausible para el grado y área indicada.
        - La temática debe ser la de la {estacion_elegida}, y esto debe ser central, no una mera contextualización.
        - Debe garantizarse que el proceso cognitivo corresponde fielmente a la descripción de la taxonomia de Bloom.
        - Evita referencias a marcas, nombres propios, lugares reales o información personal identificable.

        ENUNCIADO:
        - Formula una pregunta clara, directa, sin ambigüedades ni tecnicismos innecesarios.
        - Si utilizas negaciones, resáltalas en MAYÚSCULAS Y NEGRITA (por ejemplo: **NO ES**, **EXCEPTO**).
        - Asegúrate de que el enunciado refleje el tipo de tarea cognitiva esperado según el proceso de Bloom.

        OPCIONES DE RESPUESTA:
        - Escribe exactamente cuatro opciones (A, B, C  y D).
        - Solo una opción debe ser correcta.
        - Los distractores (respuestas incorrectas) deben estar bien diseñados: deben ser creíbles, funcionales y representar errores comunes o concepciones alternativas frecuentes.
        - No utilices fórmulas vagas como “ninguna de las anteriores” o “todas las anteriores”.

        JUSTIFICACIONES:
        {formato_justificacion}

        --- PROMPT ADICIONAL: REGLAS GENERALES DE CONSTRUCCIÓN ---
        {prompt_construccion_adicional if prompt_construccion_adicional else "No se proporcionaron prompts adicionales específicos para reglas generales de construcción."}
        ---------------------------------------------------------

        --- REGLAS ADICIONALES DEL MANUAL DE CONSTRUCCIÓN ---
        Considera y aplica estrictamente todas las directrices, ejemplos y restricciones contenidas en el siguiente manual.
        Esto es de suma importancia para la calidad y pertinencia del ítem.

        Manual de Reglas:
        {manual_reglas_texto}
        ----------------------------------------------------

        --- INFORMACIÓN ADICIONAL PROPORCIONADA POR EL USUARIO (Contexto General) ---
        {informacion_adicional_usuario if informacion_adicional_usuario else "No se proporcionó información adicional general."}
        ---------------------------------------------------------------------------
        
        --- PROMPT ADICIONAL: COSAS ESPECÍFICAS A TENER EN CUENTA ---
        {prompt_especifico_adicional if prompt_especifico_adicional else "No se proporcionaron prompts adicionales específicos para consideraciones adicionales."}
        ----------------------------------------------------------

        --- DATO CLAVE PARA LA CONSTRUCCIÓN ---
        Basado en el foco temático y el proceso cognitivo, considera el siguiente dato o idea esencial:
        "{dato_para_pregunta_foco}"

        --- INSTRUCCIONES ESPECÍFICAS DE SALIDA PARA GRÁFICO ---
        Después del bloque de JUSTIFICACIONES, incluye la siguiente información para indicar si el ítem necesita un gráfico y cómo sería:
        GRAFICO_NECESARIO: [SÍ/NO]
        DESCRIPCION_GRAFICO: [Si GRAFICO_NECESARIO es SÍ, proporciona una descripción MUY DETALLADA del gráfico. Incluye: tipo de gráfico (ej. barras, líneas, circular, diagrama de flujo, imagen de un objeto), datos o rangos de valores, etiquetas de ejes, elementos clave, propósito del gráfico y cómo se relaciona con la pregunta. Si es NO, escribe N/A.]

        --- FORMATO ESPERADO DE SALIDA ---
        PREGUNTA: [Redacta aquí el enunciado de la pregunta]
        A. [Opción A]  
        B. [Opción B]  
        C. [Opción C] 
        D. [Opción D]          
        RESPUESTA CORRECTA: [Letra de la opción correcta, por ejemplo: B]
        JUSTIFICACIONES:  
        A. [Explica por qué A es incorrecta o correcta]  
        B. [Explica por qué B es incorrecta o correcta]  
        C. [Explica por qué C es incorrecta o correcta]  
        D. [Explica por qué D es incorrecta o correcta]  
        GRAFICO_NECESARIO: [SÍ/NO]
        DESCRIPCION_GRAFICO: [Descripción detallada o N/A]
        """
        
        # Si no es el primer intento, añade las observaciones de auditoría para refinamiento
        if attempt > 1:
            prompt_content_for_llm += f"""
            --- RETROALIMENTACIÓN DE AUDITORÍA PARA REFINAMIENTO ---
            El ítem anterior no cumplió con todos los criterios. Por favor, revisa las siguientes observaciones y mejora el ítem para abordarlas.
            Observaciones del Auditor:
            {audit_observations}
            ---------------------------------------------------
            """
            # Agrega el ítem anterior para que el LLM lo pueda reformular
            prompt_content_for_llm += f"""
            --- ÍTEM ANTERIOR A REFINAR ---
            {current_item_text}
            -------------------------------
            """
        
        # Añade el prompt de feedback del usuario si existe
        prompt_content_for_llm += prompt_con_feedback
        
        # Guardar el prompt completo del generador antes de enviarlo
        full_generation_prompt = prompt_content_for_llm

        try:
            with st.spinner(f"Generando contenido con IA ({gen_model_type} - {gen_model_name}, Intento {attempt})..."):
                full_llm_response = generar_texto_con_llm(gen_model_type, gen_model_name, prompt_content_for_llm)
                
                if full_llm_response is None: # Si hubo un error en la generación con LLM
                    # st.error(f"Fallo en la generación de texto con {gen_model_type} ({gen_model_name}).") # Comentado para no saturar si son muchos ítems
                    auditoria_status = "❌ RECHAZADO (Error de Generación)"
                    audit_observations = "El modelo de generación no pudo producir una respuesta válida."
                    break # Salir del bucle de refinamiento
                    
                # --- Parsear la respuesta para extraer el ítem y la información del gráfico ---
                item_and_graphic_match = re.search(r"(PREGUNTA:.*?)(GRAFICO_NECESARIO:\s*(SÍ|NO).*?DESCRIPCION_GRAFICO:.*)", full_llm_response, re.DOTALL)
                
                if item_and_graphic_match:
                    current_item_text = item_and_graphic_match.group(1).strip()
                    grafico_info_block = item_and_graphic_match.group(2).strip()
                    
                    grafico_necesario_match = re.search(r"GRAFICO_NECESARIO:\s*(SÍ|NO)", grafico_info_block)
                    if grafico_necesario_match:
                        grafico_necesario = grafico_necesario_match.group(1).strip()

                    descripcion_grafico_match = re.search(r"DESCRIPCION_GRAFICO:\s*(.*)", grafico_info_block, re.DOTALL)
                    if descripcion_grafico_match:
                        descripcion_grafico = descripcion_grafico_match.group(1).strip()
                        if descripcion_grafico.upper() == 'N/A':
                            descripcion_grafico = ""
                else:
                    current_item_text = full_llm_response
                    grafico_necesario = "NO"
                    descripcion_grafico = ""
                    st.warning("No se pudo parsear el formato de gráfico de la respuesta. Asumiendo que no requiere gráfico.")
                
            with st.spinner(f"Auditando ítem ({audit_model_type} - {audit_model_name}, Intento {attempt})..."):
                auditoria_resultado, full_auditor_prompt = auditar_item_con_llm( # Recibe también el prompt del auditor
                    audit_model_type, audit_model_name,
                    item_generado=current_item_text,
                    grado=grado_elegido, area=area_elegida, asignatura=asignatura_seleccionada, estacion=estacion_elegida,
                    proceso_cognitivo=proceso_cognitivo_seleccionado, nanohabilidad=nanohabilidad_seleccionada,
                    microhabilidad=microhabilidad_elegida, competencia_nanohabilidad=competencia_nanohabilidad_elegida,
                    contexto_educativo=contexto_educativo, manual_reglas_texto=manual_reglas_texto,
                    descripcion_bloom=descripcion_bloom,
                    grafico_necesario=grafico_necesario,
                    descripcion_grafico=descripcion_grafico,
                    prompt_auditor_adicional=prompt_auditor_adicional # Pasa el prompt adicional del auditor
                )
                if auditoria_resultado is None: # Si hubo un error en la auditoría con LLM
                    # st.error(f"Fallo en la auditoría con {audit_model_type} ({audit_model_name}).") # Comentado
                    auditoria_status = "❌ RECHAZADO (Error de Auditoría)"
                    audit_observations = "El modelo de auditoría no pudo producir una respuesta válida."
                    break # Salir del bucle de refinamiento

                # --- Extraer DICTAMEN FINAL de forma más robusta ---
                dictamen_final_match = re.search(r"DICTAMEN FINAL:\s*\[(.*?)]", auditoria_resultado, re.DOTALL)
                if dictamen_final_match:
                    auditoria_status = dictamen_final_match.group(1).strip()
                else:
                    auditoria_status = "❌ RECHAZADO (no se pudo extraer dictamen)"
                
                observaciones_start = auditoria_resultado.find("OBSERVACIONES FINALES:")
                if observaciones_start != -1:
                    audit_observations = auditoria_resultado[observaciones_start + len("OBSERVACIONES FINALES:"):].strip()
                else:
                    audit_observations = "No se pudieron extraer observaciones específicas del auditor. Posiblemente un error de formato en la respuesta del auditor."
                
            # Guardar los datos del ítem, incluyendo el estado final de la auditoría y observaciones
            item_final_data = {
                "item_text": current_item_text,
                "classification": classification_details,
                "grafico_necesario": grafico_necesario,
                "descripcion_grafico": descripcion_grafico,
                "final_audit_status": auditoria_status, 
                "final_audit_observations": audit_observations,
                "generation_prompt_used": full_generation_prompt, # Guarda el prompt exacto usado por el generador
                "auditor_prompt_used": full_auditor_prompt
            }

            if auditoria_status == "✅ CUMPLE TOTALMENTE":
                break # Sale del ciclo de refinamiento si es aprobado
            else:
                pass # Solo se registra el estado, no se muestra advertencia por cada intento en bucle masivo

        except Exception as e:
            audit_observations = f"Error técnico durante la generación: {e}. Por favor, corrige este problema."
            auditoria_status = "❌ RECHAZADO (error técnico)"  
            item_final_data = {
                "item_text": current_item_text if current_item_text else "No se pudo generar el ítem debido a un error técnico.",
                "classification": classification_details,
                "grafico_necesario": "NO",
                "descripcion_grafico": "",
                "final_audit_status": auditoria_status,
                "final_audit_observations": audit_observations,
                "generation_prompt_used": full_generation_prompt,
                "auditor_prompt_used": full_auditor_prompt
            }
            break # Salir del ciclo si hay un error técnico grave

    if item_final_data is None:  
        return None # Retorna None si no se logró generar nada en absoluto.

    return item_final_data # Retorna el diccionario del ítem procesado

# --- Función para exportar preguntas a un documento Word ---
def exportar_a_word(preguntas_procesadas_list):
    """
    Exporta una lista de preguntas procesadas a un documento de Word (.docx) en memoria,
    incluyendo sus detalles de clasificación, la descripción del gráfico si aplica,
    y el dictamen final de la auditoría.
    Returns: BytesIO object of the document.
    """
    doc = docx.Document()
    
    doc.add_heading('Preguntas Generadas y Auditadas', level=1)
    doc.add_paragraph('Este documento contiene los ítems generados por el sistema de IA y sus resultados de auditoría.')
    doc.add_paragraph('') # Espacio en blanco

    if not preguntas_procesadas_list:
        doc.add_paragraph('No se procesaron ítems para este informe.')

    for i, item_data in enumerate(preguntas_procesadas_list):
        pregunta_texto = item_data["item_text"]
        classification = item_data["classification"]
        grafico_necesario = item_data.get("grafico_necesario", "NO")
        descripcion_grafico = item_data.get("descripcion_grafico", "")
        final_audit_status = item_data.get("final_audit_status", "N/A")
        final_audit_observations = item_data.get("final_audit_observations", "No hay observaciones finales de auditoría.")

        doc.add_heading(f'Ítem #{i+1}', level=2)
        
        # Añadir detalles de clasificación
        doc.add_paragraph('--- Clasificación del Ítem ---') 
        for key, value in classification.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{key}: ")
            run.bold = True
            p.add_run(str(value)) 

        doc.add_paragraph('') 
        
        # Añadir el texto de la pregunta y su formato
        lines = pregunta_texto.split('\n')
        for line in lines:
            line = line.strip() 
            if not line: 
                continue

            if line.startswith("PREGUNTA:"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = docx.shared.Pt(12) 
            elif line.startswith("A.") or line.startswith("B.") or line.startswith("C.")or line.startswith("D."):
                p = doc.add_paragraph(line)
                p.paragraph_format.left_indent = docx.shared.Inches(0.5) 
            elif line.startswith("RESPUESTA CORRECTA:"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            elif line.startswith("JUSTIFICACIONES:"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            elif line.startswith("GRAFICO_NECESARIO:") or line.startswith("DESCRIPCION_GRAFICO:"):
                continue 
            elif line.startswith("VALIDACIÓN DE CRITERIOS:") or line.startswith("DICTAMEN FINAL:") or line.startswith("OBSERVACIONES FINALES:"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            elif line.startswith("✅") or line.startswith("⚠️") or line.startswith("❌"):
                p = doc.add_paragraph(line)
                p.paragraph_format.left_indent = docx.shared.Inches(0.25) 
            else:
                doc.add_paragraph(line)
        
        # Añadir descripción del gráfico si es necesario
        if grafico_necesario == "SÍ" and descripcion_grafico:
            doc.add_paragraph('')
            p = doc.add_paragraph()
            run = p.add_run("--- Gráfico Sugerido ---")
            run.bold = True
            doc.add_paragraph(f"**Tipo y Descripción del Gráfico:** {descripcion_grafico}")
            doc.add_paragraph('') 

        # Añadir el dictamen final y las observaciones de la auditoría para CADA ítem
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run("--- Resultado Final de Auditoría ---")
        run.bold = True
        doc.add_paragraph(f"**DICTAMEN FINAL:** {final_audit_status}")
        doc.add_paragraph(f"**OBSERVACIONES FINALES:** {final_audit_observations}")
        doc.add_paragraph('') 

        doc.add_page_break() 

    # Guardar el documento en un buffer en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0) 
    return buffer

# --- Interfaz de Usuario Principal de Streamlit ---
st.title("📚 Generador y Auditor de ítems para el proyecto SUMUN 🧠")
st.markdown("Esta aplicación genera ítems de selección múltiple y audita su calidad, permitiéndote guiar a la IA con prompts adicionales.")

# Sección de Carga de Archivos Global (Excel y PDF)
st.sidebar.header("Carga de Archivos Global")
uploaded_excel_file = st.sidebar.file_uploader("Sube tu archivo Excel (ESTRUCTURA_TOTAL.xlsx)", type=["xlsx"])
uploaded_pdf_file = st.sidebar.file_uploader("Sube tu archivo PDF (Manual_construccion_pruebas_IMProve.pdf)", type=["pdf"])

df_datos = None
manual_reglas_texto = ""

if uploaded_excel_file:
    df_datos = leer_excel_cargado(uploaded_excel_file)

if uploaded_pdf_file:
    manual_reglas_texto = leer_pdf_cargado(uploaded_pdf_file)
    max_manual_length = 15000 
    if len(manual_reglas_texto) > max_manual_length:
        st.sidebar.warning(f"Manual es demasiado largo ({len(manual_reglas_texto)} caracteres). Truncando a {max_manual_length} caracteres para la IA.")
        manual_reglas_texto = manual_reglas_texto[:max_manual_length] # Corregido typo
    st.sidebar.info(f"Manual de reglas cargado. Longitud final: {len(manual_reglas_texto)} caracteres.")

# --- Lógica principal de Generación y Auditoría de Ítems ---
st.header("Generación y Auditoría de Ítems.")
st.markdown("Define los criterios del ítem y utiliza modelos de IA para generarlo y validarlo.")

if df_datos is None:
    st.info("Para comenzar, por favor sube tu archivo **Excel** con la estructura de datos en la barra lateral.")
elif not (gemini_config_ok or openai_config_ok):
    st.info("Para usar los modelos de IA, por favor ingresa al menos una **API Key de Gemini o OpenAI** en la barra lateral.")
else:
    st.subheader("Selecciona los Criterios para la Generación")

    # Obtener valores únicos para cada columna para los selectbox
    all_grades = df_datos['GRADO'].dropna().unique().tolist()
    grado_seleccionado = st.selectbox("Grado", sorted(all_grades), key="grado_sel")

    # Filtrar el DataFrame según la selección del grado
    df_filtrado_grado = df_datos[df_datos['GRADO'].astype(str).str.upper() == str(grado_seleccionado).upper()]
    all_areas = df_filtrado_grado['ÁREA'].dropna().unique().tolist()
    area_seleccionada = st.selectbox("Área", sorted(all_areas), key="area_sel")

    # Filtrar según la selección del área
    df_filtrado_area = df_filtrado_grado[df_filtrado_grado['ÁREA'].astype(str).str.upper() == str(area_seleccionada).upper()]
    all_asignaturas = df_filtrado_area['ASIGNATURA'].dropna().unique().tolist()
    asignatura_seleccionada = st.selectbox("Asignatura", sorted(all_asignaturas), key="asignatura_sel")

    # Filtrar según la selección de asignatura
    df_filtrado_asignatura = df_filtrado_area[df_filtrado_area['ASIGNATURA'].astype(str).str.upper() == str(asignatura_seleccionada).upper()]
    all_estaciones = df_filtrado_asignatura['ESTACIÓN'].dropna().unique().tolist()
    estacion_seleccionada = st.selectbox("Estación", sorted(all_estaciones), key="estacion_sel")

    # Filtrar según la selección de estación
    df_filtrado_estacion = df_filtrado_asignatura[df_filtrado_asignatura['ESTACIÓN'].astype(str).str.upper() == str(estacion_seleccionada).upper()]
    
    # --- Nueva Lógica: Generar todos los ítems de la estación o uno específico ---
    st.markdown("---")
    st.subheader("Modo de Generación de Ítems")
    generate_all_for_station = st.checkbox(
        "Generar TODOS los ítems de esta Estación (uno por cada Proceso Cognitivo asociado)",
        key="generate_all_station"
    )

    contexto_general_estacion = ""
    if generate_all_for_station:
        st.info("Has seleccionado generar ítems para todos los procesos cognitivos asociados a esta estación.")
        st.markdown("##### Configuración del Contexto General para la Estación")
        context_option = st.radio(
            "¿Cómo deseas establecer el contexto general para todos los ítems de esta estación?",
            ("Yo quiero dar una idea del contexto general", "Quiero que el contexto general sea generado por la IA"),
            key="context_gen_option",
            horizontal=True
        )

        if context_option == "Yo quiero dar una idea del contexto general":
            contexto_general_estacion = st.text_area(
                "Escribe tu idea o directriz para el contexto general de la estación (ej: 'Un viaje escolar a un ecosistema de páramo'):",
                height=150, key="user_context_idea"
            )
            if not contexto_general_estacion.strip():
                st.warning("Por favor, introduce una idea para el contexto general o selecciona que la IA lo genere.")
        else:
            st.info("La IA generará un contexto general apropiado para esta estación.")
            # No se necesita input del usuario, el LLM lo generará internamente.

    # Filtrar para proceso cognitivo y nanohabilidad SOLO SI NO SE GENERAN TODOS LOS DE LA ESTACIÓN
    proceso_cognitivo_seleccionado = None
    nanohabilidad_seleccionada = None
    df_item_seleccionado = None

    if not generate_all_for_station:
        all_procesos = df_filtrado_estacion['PROCESO COGNITIVO'].dropna().unique().tolist()
        proceso_cognitivo_seleccionado = st.selectbox("Proceso Cognitivo", sorted(all_procesos), key="proceso_sel")

        df_filtrado_proceso = df_filtrado_estacion[df_filtrado_estacion['PROCESO COGNITIVO'].astype(str).str.upper() == str(proceso_cognitivo_seleccionado).upper()]
        all_nanohabilidades = df_filtrado_proceso['NANOHABILIDAD'].dropna().unique().tolist()
        nanohabilidad_seleccionada = st.selectbox("Nanohabilidad", sorted(all_nanohabilidades), key="nanohabilidad_sel")

        # Se filtra el DataFrame final para un solo ítem
        df_item_seleccionado = df_filtrado_proceso[df_filtrado_proceso['NANOHABILIDAD'].astype(str).str.upper() == str(nanohabilidad_seleccionada).upper()]
        
        if df_item_seleccionado.empty:
            st.warning("No se encontraron datos en el Excel para la combinación de criterios seleccionada. Por favor, ajusta tus filtros.")
            # st.stop() # No detener la ejecución, solo mostrar el mensaje
    else: # Si se generan todos, necesitamos las filas de la estación completa
        df_item_seleccionado = df_filtrado_estacion.copy() # Copiamos todas las filas de la estación


    if df_item_seleccionado.empty:
        st.error("No hay datos válidos para generar ítems con los filtros actuales. Por favor, revisa tus selecciones o el archivo Excel.")
    else: # Solo mostrar opciones de generación si hay datos válidos
        # --- Información Adicional del Usuario (contexto general para el ítem, si no es una estación completa) ---
        if not generate_all_for_station: # Solo mostrar si es generación de ítem único, el otro ya tiene contexto general
            st.subheader("Contexto Adicional para el Ítem (Opcional - solo para ítem individual)")
            opcion_info_adicional = st.radio(
                "¿Deseas proporcionar alguna información o contexto adicional para este ítem individual?",
                ("No", "Sí"),
                key="info_ad_radio",
                horizontal=True
            )
            informacion_adicional_usuario = ""
            if opcion_info_adicional == "Sí":
                informacion_adicional_usuario = st.text_area("Escribe aquí el contexto o la información que consideres relevante para la creación del ítem:", key="info_ad_text")
        else:
            informacion_adicional_usuario = "" # No se usa si se genera por estación

        # --- Nueva Sección: Usar Prompts Adicionales ---
        st.markdown("---")
        st.subheader("Personaliza con Prompts Adicionales (Opcional)")
        use_additional_prompts = st.checkbox("Activar Prompts Adicionales", help="Si activas esto, podrás añadir instrucciones específicas para el generador y/o el auditor.")
        
        # Inicializar prompts adicionales
        prompt_bloom_adicional = ""
        prompt_construccion_adicional = ""
        prompt_especifico_adicional = ""
        prompt_auditor_adicional = ""

        if use_additional_prompts:
            col_gen_prompt, col_audit_prompt = st.columns(2)
            with col_gen_prompt:
                st.markdown("##### Prompts para el **Generador** de Ítems")
                
                # Opción 1: Prompts acerca de procesos cognitivos/Taxonomía de Bloom
                use_bloom_prompt = st.checkbox("Prompts acerca de Procesos Cognitivos / Taxonomía de Bloom", key="chk_bloom_prompt")
                if use_bloom_prompt:
                    prompt_bloom_adicional = st.text_area(
                        "Instrucciones para el generador sobre cómo aplicar la Taxonomía de Bloom (ej: 'El ítem debe requerir un análisis profundo de causas y efectos', 'Asegúrate de que el proceso cognitivo sea estrictamente de 'RECORDAR' y no de 'COMPRENDER'):", 
                        height=100, 
                        key="gen_bloom_prompt_text"
                    )

                # Opción 2: Prompts acerca de cosas generales de construcción
                use_construccion_prompt = st.checkbox("Prompts acerca de Reglas Generales de Construcción", key="chk_construccion_prompt")
                if use_construccion_prompt:
                    prompt_construccion_adicional = st.text_area(
                        "Instrucciones para el generador sobre el formato general o estilo (ej: 'Evita frases pasivas en el enunciado', 'Las opciones deben ser de longitud similar'):", 
                        height=100, 
                        key="gen_construccion_prompt_text"
                    )

                # Opción 3: Prompts acerca de cosas específicas a tener en cuenta adicionales
                use_especifico_prompt = st.checkbox("Prompts acerca de Consideraciones Específicas Adicionales", key="chk_especifico_prompt")
                if use_especifico_prompt:
                    prompt_especifico_adicional = st.text_area(
                        "Instrucciones muy específicas o de último minuto para el generador (ej: 'El contexto debe mencionar una actividad deportiva', 'Incluye un personaje llamado 'Sofía' en el enunciado'):", 
                        height=100, 
                        key="gen_especifico_prompt_text"
                    )

            with col_audit_prompt:
                st.markdown("##### Prompts para el **Auditor** de Ítems")
                prompt_auditor_adicional = st.text_area(
                    "Instrucciones específicas para que la IA audite el ítem (ej: 'Verifica la coherencia en la numeración', 'Asegúrate que no haya ambigüedad en las opciones'):", 
                    height=200, 
                    key="audit_prompt_add"
                )

        st.markdown("---")
        # --- Selección de Modelos para Generación/Auditoría ---
        st.subheader("Configuración de Modelos de IA")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Modelo para Generación de Ítems**")
            gen_model_type = st.radio("Tipo de Modelo (Generación)", ["Gemini", "GPT"], key="gen_model_type", index=0 if gemini_config_ok else 1) 
            gen_model_name = ""
            if gen_model_type == "Gemini":
                gen_model_name = st.selectbox("Nombre del Modelo Gemini (Generación)", ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.5-flash-lite", "gemini-1.5-pro", "gemini-1.5-flash"], key="gen_gemini_name")
            else: 
                gen_model_name = st.selectbox("Nombre del Modelo GPT (Generación)", ["gpt-4o", "gpt-4-turbo", "gpt-3.5-turbo"], key="gen_gpt_name")
        
        with col2:
            st.markdown("**Modelo para Auditoría de Ítems**")
            audit_model_type = st.radio("Tipo de Modelo (Auditoría)", ["Gemini", "GPT"], key="audit_model_type", index=0 if gemini_config_ok else 1)
            audit_model_name = ""
            if audit_model_type == "Gemini":
                audit_model_name = st.selectbox("Nombre del Modelo Gemini (Auditoría)", ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.5-flash-lite", "gemini-1.5-pro", "gemini-1.5-flash"], key="audit_gemini_name")
            else: 
                audit_model_name = st.selectbox("Nombre del Modelo GPT (Auditoría)", ["gpt-4o", "gpt-4-turbo", "gpt-3.5-turbo"], key="audit_gpt_name")

        # --- Botón para Generar y Auditar ---
        if st.button("Generar y Auditar Ítem(s)"): # Texto del botón actualizado
            if (gen_model_type == "Gemini" and not gemini_config_ok) or (gen_model_type == "GPT" and not openai_config_ok):
                st.error(f"Por favor, configura la API Key para el modelo de generación ({gen_model_type}).")
            elif (audit_model_type == "Gemini" and not gemini_config_ok) or (audit_model_type == "GPT" and not openai_config_ok):
                st.error(f"Por favor, configura la API Key para el modelo de auditoría ({audit_model_type}).")
            elif generate_all_for_station and context_option == "Yo quiero dar una idea del contexto general" and not contexto_general_estacion.strip():
                st.error("Por favor, introduce tu idea para el contexto general o selecciona que la IA lo genere.")
            else:
                st.markdown("---")
                st.info("Iniciando generación y auditoría del(los) ítem(s). Esto puede tardar unos momentos...")

                criterios_para_preguntas = {
                    "tipo_pregunta": "opción múltiple con 4 opciones",  
                    "dificultad": "media", 
                    "num_preguntas": 1,   
                    "contexto_educativo": "estudiantes de preparatoria (bachillerato)", 
                    "formato_justificacion": """
                         • Justificación correcta: debe explicar el razonamiento o proceso cognitivo (NO por descarte).
                         • Justificaciones incorrectas: deben redactarse como: “El estudiante podría escoger la opción X porque… Sin embargo, esto es incorrecto porque…”
                    """
                }
                
                # Lista para almacenar todos los ítems procesados
                processed_items_list = []

                if generate_all_for_station:
                    st.subheader(f"Generando ítems para la Estación: {estacion_seleccionada}")
                    
                    # Asegurar que se procesa una copia para no modificar el df original
                    unique_procesos = df_item_seleccionado[['PROCESO COGNITIVO', 'NANOHABILIDAD', 'MICROHABILIDAD', 'COMPETENCIA NANOHABILIDAD']].drop_duplicates().to_dict('records')
                    
                    if not unique_procesos:
                        st.warning(f"No se encontraron procesos cognitivos únicos para la estación '{estacion_seleccionada}'. No se generarán ítems.")
                    else:
                        progress_bar_text = st.empty()
                        progress_bar = st.progress(0)
                        
                        for i, item_spec_row in enumerate(unique_procesos):
                            # Construir una 'fila_datos' simulada para cada proceso cognitivo para la función de generación
                            current_fila_datos = {
                                'GRADO': grado_seleccionado,
                                'ÁREA': area_seleccionada,
                                'ASIGNATURA': asignatura_seleccionada,
                                'ESTACIÓN': estacion_seleccionada,
                                'PROCESO COGNITIVO': item_spec_row['PROCESO COGNITIVO'],
                                'NANOHABILIDAD': item_spec_row['NANOHABILIDAD'],
                                'MICROHABILIDAD': item_spec_row['MICROHABILIDAD'],
                                'COMPETENCIA NANOHABILIDAD': item_spec_row['COMPETENCIA NANOHABILIDAD']
                            }
                            
                            progress_text_msg = f"Procesando ítem {i+1} de {len(unique_procesos)}: {item_spec_row['PROCESO COGNITIVO']} - {item_spec_row['NANOHABILIDAD']}"
                            progress_bar_text.text(progress_text_msg)
                            progress_bar.progress((i + 1) / len(unique_procesos))

                            st.write(f"**Generando para:** {item_spec_row['PROCESO COGNITIVO']} - {item_spec_row['NANOHABILIDAD']}")

                            item_data = generar_pregunta_con_seleccion(
                                gen_model_type, gen_model_name, audit_model_type, audit_model_name, 
                                fila_datos=current_fila_datos,   
                                criterios_generacion=criterios_para_preguntas,
                                manual_reglas_texto=manual_reglas_texto,
                                informacion_adicional_usuario=informacion_adicional_usuario,
                                prompt_bloom_adicional=prompt_bloom_adicional, 
                                prompt_construccion_adicional=prompt_construccion_adicional, 
                                prompt_especifico_adicional=prompt_especifico_adicional, 
                                prompt_auditor_adicional=prompt_auditor_adicional,
                                contexto_general_estacion=contexto_general_estacion if context_option == "Yo quiero dar una idea del contexto general" else "" # Pasa el contexto definido por el usuario si aplica
                            )
                            if item_data:
                                processed_items_list.append(item_data)
                            st.markdown("---") # Separador visual entre ítems generados
                        
                        progress_bar_text.text("Todos los ítems han sido procesados.")
                        progress_bar.progress(1.0)
                        st.success(f"Se han procesado {len(processed_items_list)} ítems para la estación '{estacion_seleccionada}'.")
                        
                        # Guardar la lista completa para su posterior revisión
                        st.session_state['processed_items_list_for_review'] = processed_items_list
                        if processed_items_list:
                            st.session_state['current_review_index'] = 0
                            st.session_state['awaiting_review'] = True
                else: # Generación de un solo ítem
                    st.subheader(f"Generando ítem individual para: {proceso_cognitivo_seleccionado} - {nanohabilidad_seleccionada}")
                    item_data = generar_pregunta_con_seleccion(
                        gen_model_type, gen_model_name, audit_model_type, audit_model_name, 
                        fila_datos=df_item_seleccionado.iloc[0],   
                        criterios_generacion=criterios_para_preguntas,
                        manual_reglas_texto=manual_reglas_texto,
                        informacion_adicional_usuario=informacion_adicional_usuario,
                        prompt_bloom_adicional=prompt_bloom_adicional, 
                        prompt_construccion_adicional=prompt_construccion_adicional, 
                        prompt_especifico_adicional=prompt_especifico_adicional, 
                        prompt_auditor_adicional=prompt_auditor_adicional,
                        contexto_general_estacion="" # No hay contexto de estación si es ítem individual
                    )
                    if item_data:
                        st.session_state['processed_items_list_for_review'] = [item_data]
                        st.session_state['current_review_index'] = 0
                        st.session_state['awaiting_review'] = True
                    else:
                        st.error("No se pudo generar el ítem bajo las condiciones seleccionadas.")

# --- Lógica para mostrar la interfaz de revisión ---
if 'awaiting_review' in st.session_state and st.session_state['awaiting_review']:
    
    if 'approved_items' not in st.session_state:
        st.session_state['approved_items'] = []
    
    current_index = st.session_state.get('current_review_index', 0)
    items_to_review = st.session_state.get('processed_items_list_for_review', [])
    
    if current_index >= len(items_to_review):
        st.success("¡Has revisado todos los ítems! Ahora puedes descargarlos en la sección de 'Exportar Resultados'.")
        st.session_state['awaiting_review'] = False
        st.session_state['current_review_index'] = 0
        del st.session_state['processed_items_list_for_review']
        st.rerun()
        
    item_to_review = items_to_review[current_index]

    st.markdown("---")
    st.header(f"Revisión de Ítem ({current_index + 1} de {len(items_to_review)})")
    st.info(f"Dictamen de la Auditoría Inicial: **{item_to_review['final_audit_status']}**")
    
    st.markdown("---")
    st.markdown("### Ítem Generado:")
    st.markdown(item_to_review['item_text'])
    
    st.markdown("---")
    st.markdown("### Observaciones de la Auditoría:")
    st.markdown(item_to_review['final_audit_observations'])
    
    col_aprob, col_rechazo = st.columns(2)

    with col_aprob:
        if st.button("✅ Aprobar y Siguiente"):
            st.session_state['approved_items'].append(item_to_review)
            st.session_state['current_review_index'] += 1
            st.rerun()

    with col_rechazo:
        if st.button("❌ Rechazar y Reintentar"):
            st.session_state['show_feedback_form'] = True
            
    if st.session_state.get('show_feedback_form', False):
        st.markdown("---")
        st.markdown("#### Por favor, proporciona tus observaciones para refinar el ítem:")
        
        feedback_enunciado = st.text_area("1. Observaciones del enunciado/contexto:", key="feedback_enunciado")
        feedback_opciones = st.text_area("2. Observaciones de las opciones de respuesta:", key="feedback_opciones")
        
        # Aquí construimos el prompt de feedback completo
        feedback_completo = ""
        if feedback_enunciado:
            feedback_completo += f"Observaciones sobre el enunciado/contexto: {feedback_enunciado}\n"
        if feedback_opciones:
            feedback_completo += f"Observaciones sobre las opciones de respuesta: {feedback_opciones}\n"

        if st.button("🔄 Refinar con estas Observaciones"):
            st.info("Re-generando el ítem con tu feedback...")
            
            # Recuperar la fila de datos original
            original_fila_datos = item_to_review['classification']
            
            # Volver a llamar a la función de generación con el feedback del usuario
            refined_item_data = generar_pregunta_con_seleccion(
                gen_model_type, gen_model_name, audit_model_type, audit_model_name,
                fila_datos=original_fila_datos,
                criterios_generacion={"tipo_pregunta": "opción múltiple con 4 opciones", "dificultad": "media", "num_preguntas": 1, "contexto_educativo": "estudiantes de preparatoria (bachillerato)"},
                manual_reglas_texto=manual_reglas_texto,
                informacion_adicional_usuario=informacion_adicional_usuario,
                prompt_bloom_adicional=prompt_bloom_adicional,
                prompt_construccion_adicional=prompt_construccion_adicional,
                prompt_especifico_adicional=prompt_especifico_adicional,
                prompt_auditor_adicional=prompt_auditor_adicional,
                contexto_general_estacion=contexto_general_estacion,
                feedback_usuario=feedback_completo # Pasar el feedback al generador
            )
            
            if refined_item_data:
                # Reemplazar el ítem actual con el refinado
                st.session_state['processed_items_list_for_review'][current_index] = refined_item_data
                st.session_state['show_feedback_form'] = False
                st.success("¡Ítem refinado exitosamente! Por favor, revísalo de nuevo.")
                st.rerun()
            else:
                st.error("Fallo al refinar el ítem. Intenta de nuevo o ajusta tu feedback.")


# --- Sección de Exportación a Word y descarga de Prompts (siempre visible al final de esta sección) ---
st.header("Exportar Resultados")

if 'approved_items' in st.session_state and st.session_state['approved_items']:
    num_items_processed = len(st.session_state['approved_items'])
    st.write(f"Hay **{num_items_processed}** ítem(s) aprobado(s) y disponible(s) para exportar.")
    
    # Exportar a Word
    nombre_archivo_word = st.text_input("Ingresa el nombre deseado para el archivo Word (sin la extensión .docx):", 
                                        f"items_{estacion_seleccionada.replace(' ', '_')}_{grado_seleccionado}", 
                                        key="word_filename")
    if nombre_archivo_word:
        word_buffer = exportar_a_word(st.session_state['approved_items'])
        st.download_button(
            label="Descargar Ítem(s) en Documento Word",
            data=word_buffer,
            file_name=f"{nombre_archivo_word}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.info("Haz clic para descargar el archivo Word con el(los) ítem(s) y su(s) auditoría(s).")
    else:
        st.warning("Por favor, ingresa un nombre para el archivo Word.")

    # Descargar Prompts Utilizados
    st.markdown("---")
    st.subheader("Descargar Prompts Utilizados")
    st.info("Puedes descargar un archivo TXT con los prompts completos que se enviaron a los modelos de IA para este(os) ítem(s).")
    
    # Construir el contenido del TXT con ambos prompts para todos los ítems
    combined_prompts_content = ""
    for i, item_data in enumerate(st.session_state['approved_items']):
        combined_prompts_content += f"--- PROMPT DETALLADO PARA ÍTEM #{i+1} ---\n"
        combined_prompts_content += f"**Clasificación:** Grado: {item_data['classification']['Grado']}, Área: {item_data['classification']['Área']}, Asignatura: {item_data['classification']['Asignatura']}, Estación: {item_data['classification']['Estación']}, Proceso Cognitivo: {item_data['classification']['Proceso Cognitivo']}, Nanohabilidad: {item_data['classification']['Nanohabilidad']}\n\n"
        combined_prompts_content += f"--- PROMPT COMPLETO ENVIADO AL GENERADOR ---\n"
        combined_prompts_content += f"{item_data.get('generation_prompt_used', 'No disponible')}\n\n"
        combined_prompts_content += f"--- PROMPT COMPLETO ENVIADO AL AUDITOR ---\n"
        combined_prompts_content += f"{item_data.get('auditor_prompt_used', 'No disponible')}\n\n"
        combined_prompts_content += "="*80 + "\n\n" # Separador entre prompts de ítems
    
    prompt_download_filename = st.text_input("Nombre para el archivo TXT de prompts (sin .txt):", f"prompts_{estacion_seleccionada.replace(' ', '_')}", key="prompt_txt_filename")
    if prompt_download_filename:
        st.download_button(
            label="Descargar Prompts como TXT",
            data=combined_prompts_content.encode('utf-8'),
            file_name=f"{prompt_download_filename}.txt",
            mime="text/plain"
        )
        st.info("Haz clic para descargar el archivo TXT con los prompts detallados de todos los ítems.")
    else:
        st.warning("Ingresa un nombre para el archivo de prompts.")
else:
    st.info("No hay ítems aprobados disponibles para exportar en este momento. Genera y audita ítem(s) para que estén disponibles aquí.")
